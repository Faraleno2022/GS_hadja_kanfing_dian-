"""
Processeur de documents pour extraire le texte des fichiers uploadés
Supporte PDF, Word (docx), et fichiers texte
"""
import os
import re


def extraire_texte_pdf(fichier_path):
    """Extrait le texte d'un fichier PDF"""
    try:
        import PyPDF2
        texte = ""
        with open(fichier_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                texte += page.extract_text() or ""
        return nettoyer_texte(texte)
    except ImportError:
        # PyPDF2 non installé, essayer pdfplumber
        try:
            import pdfplumber
            texte = ""
            with pdfplumber.open(fichier_path) as pdf:
                for page in pdf.pages:
                    texte += page.extract_text() or ""
            return nettoyer_texte(texte)
        except ImportError:
            return "[Erreur: Installez PyPDF2 ou pdfplumber pour lire les PDF]"
    except Exception as e:
        return f"[Erreur lors de la lecture du PDF: {str(e)}]"


def extraire_texte_docx(fichier_path):
    """Extrait le texte d'un fichier Word (.docx)"""
    try:
        from docx import Document
        doc = Document(fichier_path)
        texte = ""
        for para in doc.paragraphs:
            texte += para.text + "\n"
        # Extraire aussi le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texte += cell.text + " "
                texte += "\n"
        return nettoyer_texte(texte)
    except ImportError:
        return "[Erreur: Installez python-docx pour lire les fichiers Word]"
    except Exception as e:
        return f"[Erreur lors de la lecture du fichier Word: {str(e)}]"


def extraire_texte_txt(fichier_path):
    """Extrait le texte d'un fichier texte"""
    try:
        # Essayer différents encodages
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(fichier_path, 'r', encoding=encoding) as f:
                    return nettoyer_texte(f.read())
            except UnicodeDecodeError:
                continue
        return "[Erreur: Impossible de décoder le fichier texte]"
    except Exception as e:
        return f"[Erreur lors de la lecture du fichier texte: {str(e)}]"


def nettoyer_texte(texte):
    """Nettoie le texte extrait"""
    if not texte:
        return ""
    
    # Supprimer les caractères de contrôle
    texte = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', texte)
    
    # Remplacer les retours à la ligne multiples
    texte = re.sub(r'\n{3,}', '\n\n', texte)
    
    # Supprimer les espaces multiples
    texte = re.sub(r' {2,}', ' ', texte)
    
    # Supprimer les espaces en début/fin de ligne
    lignes = [ligne.strip() for ligne in texte.split('\n')]
    texte = '\n'.join(lignes)
    
    return texte.strip()


def extraire_texte_document(fichier_path):
    """
    Extrait le texte d'un document selon son extension
    
    Args:
        fichier_path: Chemin vers le fichier
    
    Returns:
        Texte extrait du document
    """
    if not os.path.exists(fichier_path):
        return "[Erreur: Fichier non trouvé]"
    
    extension = os.path.splitext(fichier_path)[1].lower()
    
    extracteurs = {
        '.pdf': extraire_texte_pdf,
        '.docx': extraire_texte_docx,
        '.doc': extraire_texte_docx,  # Peut ne pas fonctionner pour les anciens .doc
        '.txt': extraire_texte_txt,
        '.md': extraire_texte_txt,
        '.rtf': extraire_texte_txt,
    }
    
    extracteur = extracteurs.get(extension)
    
    if extracteur:
        return extracteur(fichier_path)
    else:
        return f"[Format non supporté: {extension}. Formats acceptés: PDF, DOCX, TXT]"


def traiter_document_upload(document):
    """
    Traite un document uploadé et extrait son contenu
    
    Args:
        document: Instance de DocumentCours
    
    Returns:
        True si le traitement a réussi, False sinon
    """
    try:
        if document.fichier:
            fichier_path = document.fichier.path
            contenu = extraire_texte_document(fichier_path)
            
            if not contenu.startswith('[Erreur'):
                document.contenu_extrait = contenu
                document.save(update_fields=['contenu_extrait'])
                return True
            else:
                # Sauvegarder le message d'erreur pour diagnostic
                document.contenu_extrait = contenu
                document.save(update_fields=['contenu_extrait'])
                return False
        return False
    except Exception as e:
        document.contenu_extrait = f"[Erreur de traitement: {str(e)}]"
        document.save(update_fields=['contenu_extrait'])
        return False


def obtenir_formats_supportes():
    """Retourne la liste des formats de fichiers supportés"""
    return {
        'pdf': {
            'extension': '.pdf',
            'nom': 'PDF',
            'icone': '📄',
            'mime': 'application/pdf'
        },
        'docx': {
            'extension': '.docx',
            'nom': 'Word',
            'icone': '📝',
            'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        },
        'txt': {
            'extension': '.txt',
            'nom': 'Texte',
            'icone': '📃',
            'mime': 'text/plain'
        }
    }


def valider_fichier(fichier):
    """
    Valide un fichier uploadé
    
    Args:
        fichier: Fichier uploadé (UploadedFile)
    
    Returns:
        Tuple (valide, message)
    """
    if not fichier:
        return False, "Aucun fichier fourni"
    
    # Vérifier la taille (max 10 MB)
    max_size = 10 * 1024 * 1024  # 10 MB
    if fichier.size > max_size:
        return False, f"Le fichier est trop volumineux (max 10 MB)"
    
    # Vérifier l'extension
    nom_fichier = fichier.name.lower()
    extensions_valides = ['.pdf', '.docx', '.doc', '.txt', '.md']
    
    extension = os.path.splitext(nom_fichier)[1]
    if extension not in extensions_valides:
        return False, f"Format non supporté. Formats acceptés: {', '.join(extensions_valides)}"
    
    return True, "Fichier valide"
