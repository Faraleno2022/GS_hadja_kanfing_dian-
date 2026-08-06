from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "documents" / "Catalogue_MySchoolGN_2026.docx"
LOGO = ROOT / "static" / "logos" / "logo.png"
SCHOOL_PHOTO = ROOT / "static" / "images" / "carte2.jpg"

# standard_business_brief preset, with a named MySchoolGN brand-color override.
LETTER_WIDTH = 12240
LETTER_HEIGHT = 15840
CONTENT_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}

NAVY = "0E4A78"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GREEN = "278A68"
ORANGE = "F28C28"
INK = "17324D"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF6F1"
PALE_ORANGE = "FFF3E5"
BORDER = "D8E1EA"
WHITE = "FFFFFF"
BLACK = "111827"
AMBER = "A76700"
PALE_AMBER = "FFF4D6"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_image_alt_text(inline_shape, description, title=None):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", title or description)
    return inline_shape


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, spec in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), spec.get("val", "single"))
        edge.set(qn("w:sz"), str(spec.get("sz", 6)))
        edge.set(qn("w:space"), str(spec.get("space", 0)))
        edge.set(qn("w:color"), spec.get("color", BORDER))


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[edge]))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT):
    if sum(widths) != CONTENT_WIDTH:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH}, received {sum(widths)}")

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size=5):
    for row in table.rows:
        for cell in row.cells:
            spec = {"val": "single", "sz": size, "color": color}
            set_cell_border(
                cell,
                top=spec,
                bottom=spec,
                start=spec,
                end=spec,
                insideH=spec,
                insideV=spec,
            )


def set_paragraph_border(paragraph, bottom=None, top=None):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name, spec in (("top", top), ("bottom", bottom)):
        if spec is None:
            continue
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(spec.get("sz", 8)))
        edge.set(qn("w:space"), str(spec.get("space", 4)))
        edge.set(qn("w:color"), spec.get("color", BLUE))
        p_bdr.append(edge)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=False):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        r_pr.append(underline_node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=8.5, color=GRAY)


def add_tab_stop(paragraph, position_dxa, alignment="right"):
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), alignment)
    tab.set(qn("w:pos"), str(position_dxa))
    tabs.append(tab)


def define_numbering(document, ordered=False):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:hint"), "default")
    r_pr.append(fonts)
    level.append(r_pr)

    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(document, text, num_id, bold_lead=None):
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    if bold_lead and text.startswith(bold_lead):
        set_font(paragraph.add_run(bold_lead), bold=True, color=INK)
        set_font(paragraph.add_run(text[len(bold_lead) :]), color=BLACK)
    else:
        set_font(paragraph.add_run(text), color=BLACK)
    return paragraph


def add_kicker(document, text, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_font(run, size=9.5, color=GREEN, bold=True)
    run.font.letter_spacing = Pt(0.6)
    return paragraph


def add_title(document, text, size=30, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, after=8):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    set_font(paragraph.add_run(text), size=size, color=color, bold=True)
    return paragraph


def add_subtitle(document, text, align=WD_ALIGN_PARAGRAPH.LEFT, after=14):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text), size=14, color=GRAY)
    return paragraph


def add_body(document, text, after=6, color=BLACK, bold_lead=None, align=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        set_font(paragraph.add_run(bold_lead), bold=True, color=INK)
        set_font(paragraph.add_run(text[len(bold_lead) :]), color=color)
    else:
        set_font(paragraph.add_run(text), color=color)
    return paragraph


def add_section_heading(document, text, after=8):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.add_run(text)
    return paragraph


def add_small_note(document, text, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run(text), size=8.5, color=GRAY, italic=True)
    return paragraph


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_card(cell, title, body, accent=BLUE, fill="FFFFFF", detail=None):
    clear_cell(cell)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 14, "color": accent},
        bottom={"val": "single", "sz": 5, "color": BORDER},
        start={"val": "single", "sz": 5, "color": BORDER},
        end={"val": "single", "sz": 5, "color": BORDER},
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(5)
    p_title.paragraph_format.keep_with_next = True
    set_font(p_title.add_run(title), size=11.5, color=accent, bold=True)

    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(3 if detail else 0)
    p_body.paragraph_format.line_spacing = 1.05
    set_font(p_body.add_run(body), size=9.6, color=BLACK)

    if detail:
        p_detail = cell.add_paragraph()
        p_detail.paragraph_format.space_after = Pt(0)
        p_detail.paragraph_format.line_spacing = 1.0
        set_font(p_detail.add_run(detail), size=8.3, color=GRAY, italic=True)


def add_cards_grid(document, cards):
    rows = (len(cards) + 1) // 2
    table = document.add_table(rows=rows, cols=2)
    set_table_geometry(table, [4680, 4680])
    for idx, card in enumerate(cards):
        row = idx // 2
        col = idx % 2
        add_card(table.cell(row, col), **card)
    if len(cards) % 2:
        set_cell_shading(table.cell(rows - 1, 1), WHITE)
        set_cell_border(
            table.cell(rows - 1, 1),
            top={"val": "nil", "sz": 0, "color": WHITE},
            bottom={"val": "nil", "sz": 0, "color": WHITE},
            start={"val": "nil", "sz": 0, "color": WHITE},
            end={"val": "nil", "sz": 0, "color": WHITE},
        )
    return table


def add_callout(document, label, text, fill=PALE_BLUE, accent=BLUE):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    clear_cell(cell)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        start={"val": "single", "sz": 18, "color": accent},
        top={"val": "single", "sz": 4, "color": BORDER},
        bottom={"val": "single", "sz": 4, "color": BORDER},
        end={"val": "single", "sz": 4, "color": BORDER},
    )
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(f"{label}  "), size=10.5, color=accent, bold=True)
    set_font(paragraph.add_run(text), size=10.5, color=BLACK)
    return table


def add_metric_strip(document, metrics):
    table = document.add_table(rows=1, cols=len(metrics))
    widths = [CONTENT_WIDTH // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_WIDTH - sum(widths)
    set_table_geometry(table, widths)
    accents = [NAVY, GREEN, ORANGE]
    fills = [PALE_BLUE, PALE_GREEN, PALE_ORANGE]
    for idx, (label, text) in enumerate(metrics):
        cell = table.cell(0, idx)
        clear_cell(cell)
        set_cell_shading(cell, fills[idx % len(fills)])
        set_cell_border(
            cell,
            top={"val": "single", "sz": 12, "color": accents[idx % len(accents)]},
            bottom={"val": "single", "sz": 4, "color": BORDER},
            start={"val": "single", "sz": 4, "color": BORDER},
            end={"val": "single", "sz": 4, "color": BORDER},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(label.upper()), size=9, color=accents[idx % len(accents)], bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        set_font(p2.add_run(text), size=9.2, color=BLACK)
    return table


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(3)
    add_tab_stop(hp, CONTENT_WIDTH, "right")
    set_font(hp.add_run("MYSCHOOLGN"), size=8.5, color=NAVY, bold=True)
    set_font(hp.add_run("\tCATALOGUE 2026"), size=8.5, color=GRAY, bold=True)
    set_paragraph_border(hp, bottom={"sz": 5, "space": 3, "color": BORDER})

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(0)
    add_tab_stop(fp, CONTENT_WIDTH, "right")
    set_font(fp.add_run("www.myschoolgn.space"), size=8.5, color=GRAY)
    set_font(fp.add_run("\tPage "), size=8.5, color=GRAY)
    add_page_field(fp)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ffp.paragraph_format.space_before = Pt(0)
    ffp.paragraph_format.space_after = Pt(0)
    set_font(
        ffp.add_run("www.myschoolgn.space  |  +224 622 61 35 59  |  contact@myschoolgn.space"),
        size=8.7,
        color=GRAY,
    )
    return section


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "Catalogue MySchoolGN 2026"
    document.core_properties.subject = "Présentation commerciale du logiciel MySchoolGN"
    document.core_properties.author = "MySchoolGN"
    document.core_properties.keywords = "MySchoolGN, gestion scolaire, Guinée, catalogue"
    configure_styles(document)
    configure_section(document)
    bullet_num_id = define_numbering(document, ordered=False)
    ordered_num_id = define_numbering(document, ordered=True)

    # Page 1 — Cover
    p_logo = document.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(4)
    set_image_alt_text(
        p_logo.add_run().add_picture(str(LOGO), width=Inches(1.25)),
        "Logo Africa Mindset Consulting",
        "Africa Mindset Consulting",
    )

    add_kicker(document, "Catalogue de solution · Édition 2026", WD_ALIGN_PARAGRAPH.CENTER, 4)
    add_title(document, "MySchoolGN", 32, NAVY, WD_ALIGN_PARAGRAPH.CENTER, 3)
    add_subtitle(
        document,
        "Le pilotage scolaire, pédagogique et financier dans un seul environnement",
        WD_ALIGN_PARAGRAPH.CENTER,
        12,
    )

    p_image = document.add_paragraph()
    p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_image.paragraph_format.space_before = Pt(0)
    p_image.paragraph_format.space_after = Pt(10)
    set_image_alt_text(
        p_image.add_run().add_picture(str(SCHOOL_PHOTO), width=Inches(6.5)),
        "Cour d'un établissement scolaire équipé d'infrastructures modernes",
        "Établissement scolaire",
    )

    add_body(
        document,
        "Une solution conçue pour les établissements guinéens, de la maternelle à la terminale, "
        "accessible en ligne et disponible en application Windows.",
        after=10,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_metric_strip(
        document,
        [
            ("Pédagogie", "Notes, bulletins et suivi des résultats"),
            ("Administration", "Élèves, classes et documents scolaires"),
            ("Finances", "Paiements, dépenses et pilotage"),
        ],
    )
    add_small_note(
        document,
        "Centralisez l'information. Automatisez les tâches répétitives. Rapprochez l'école des familles.",
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Page 2 — Promise and benefits
    add_page_break(document)
    add_kicker(document, "Une école mieux organisée")
    add_title(document, "La gestion scolaire devient plus simple", 27, NAVY, after=6)
    add_subtitle(
        document,
        "MySchoolGN réunit les opérations essentielles de l'établissement dans un système cohérent, "
        "sécurisé et facile à consulter.",
        after=12,
    )

    add_cards_grid(
        document,
        [
            {
                "title": "CENTRALISER",
                "body": "Dossiers élèves, classes, notes, paiements, personnel et services sont regroupés dans un même espace.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "AUTOMATISER",
                "body": "Moyennes, classements, échéanciers et documents PDF sont calculés ou générés sans ressaisie inutile.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
            {
                "title": "PILOTER",
                "body": "Les tableaux de bord donnent une vision claire des effectifs, des résultats, des recettes et des impayés.",
                "accent": ORANGE,
                "fill": PALE_ORANGE,
            },
            {
                "title": "COMMUNIQUER",
                "body": "Parents et responsables reçoivent plus facilement bulletins, reçus, rapports et rappels par les canaux prévus.",
                "accent": BLUE,
                "fill": LIGHT_GRAY,
            },
        ],
    )

    add_section_heading(document, "Une réponse pour chaque acteur", after=6)
    add_list_item(
        document,
        "Direction : une vue d'ensemble pour suivre l'activité et prendre des décisions documentées.",
        bullet_num_id,
        "Direction :",
    )
    add_list_item(
        document,
        "Enseignants : une saisie structurée des notes et des bulletins calculés automatiquement.",
        bullet_num_id,
        "Enseignants :",
    )
    add_list_item(
        document,
        "Comptabilité : des échéanciers, reçus et états financiers mieux organisés.",
        bullet_num_id,
        "Comptabilité :",
    )
    add_list_item(
        document,
        "Parents : un accès simple aux résultats et à la situation financière de leur enfant.",
        bullet_num_id,
        "Parents :",
    )

    add_callout(
        document,
        "DEUX MODES D'UTILISATION",
        "Accès en ligne depuis un navigateur, ou application de bureau Windows pouvant fonctionner sans connexion Internet.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    # Page 3 — Core modules
    add_page_break(document)
    add_kicker(document, "Fonctionnalités")
    add_title(document, "Les modules essentiels", 27, NAVY, after=5)
    add_subtitle(
        document,
        "Du premier contact avec la famille jusqu'au bulletin et au suivi des paiements, "
        "MySchoolGN couvre le cycle de gestion quotidien.",
        after=10,
    )
    add_cards_grid(
        document,
        [
            {
                "title": "1. GESTION DES ÉLÈVES",
                "body": "Inscriptions, dossiers, responsables, photos, matricules automatiques, classes, transferts et nouvelle année scolaire.",
                "detail": "Imports et exports Excel, cartes scolaires et fiches d'inscription PDF.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "2. NOTES & BULLETINS",
                "body": "Saisie standard, guinéenne ou intelligente, calcul des moyennes, rangs, mentions et appréciations.",
                "detail": "Bulletins primaire, secondaire et maternelle, individuels ou par classe.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
            {
                "title": "3. FINANCES & PAIEMENTS",
                "body": "Grilles tarifaires, échéanciers, paiements partiels, remises, impayés, reçus et exports comptables.",
                "detail": "Espèces, Mobile Money, chèque ou virement selon la configuration.",
                "accent": ORANGE,
                "fill": PALE_ORANGE,
            },
            {
                "title": "4. ESPACE PARENTS",
                "body": "Consultation sécurisée des notes, classements, activités, paiements et reçus depuis un téléphone ou un ordinateur.",
                "detail": "Accès par matricule, numéro du responsable et classe.",
                "accent": BLUE,
                "fill": LIGHT_GRAY,
            },
            {
                "title": "5. RAPPORTS & TABLEAUX DE BORD",
                "body": "Synthèses pédagogiques, financières et administratives, graphiques interactifs et exports PDF ou Excel.",
                "detail": "Suivi des effectifs, résultats, recettes, dépenses et taux de recouvrement.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "6. COMMUNICATION",
                "body": "Envoi de bulletins, reçus et rappels de paiement par WhatsApp ou SMS selon les services activés.",
                "detail": "Historique et suivi des envois disponibles dans le système.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
        ],
    )
    add_callout(
        document,
        "DOCUMENTS GÉNÉRÉS",
        "Cartes scolaires, bulletins, livrets, reçus, notes de rappel, certificats, fiches de paie et rapports professionnels.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Page 4 — Extended modules and implementation
    add_page_break(document)
    add_kicker(document, "Une plateforme évolutive")
    add_title(document, "Pensée pour la réalité du terrain", 27, NAVY, after=5)
    add_subtitle(
        document,
        "Les modules complémentaires permettent d'étendre MySchoolGN au rythme des priorités de l'établissement.",
        after=10,
    )
    add_cards_grid(
        document,
        [
            {
                "title": "RESSOURCES HUMAINES",
                "body": "Fiches enseignants, affectations, pointage, calcul des salaires et bulletins de paie PDF.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "DÉPENSES & FOURNISSEURS",
                "body": "Catégories, pièces justificatives, workflow de validation, budget et suivi des engagements.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
            {
                "title": "TRANSPORT & CANTINE",
                "body": "Abonnements, itinéraires, échéances, repas, restrictions alimentaires et présences.",
                "accent": ORANGE,
                "fill": PALE_ORANGE,
            },
            {
                "title": "BIBLIOTHÈQUE & LOGISTIQUE",
                "body": "Catalogue, emprunts, réservations, stocks, inventaires, biens et maintenance.",
                "accent": BLUE,
                "fill": LIGHT_GRAY,
            },
            {
                "title": "SÉCURITÉ & PERMISSIONS",
                "body": "Rôles, droits granulaires, validation des comptes, journal d'activité et accès contrôlés.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "ASSISTANT INTELLIGENT",
                "body": "Aide interactive, ressources de cours et accompagnement des utilisateurs selon les fonctions activées.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
        ],
    )

    add_section_heading(document, "Un déploiement progressif", after=5)
    add_list_item(
        document,
        "Cadrer les besoins : cycles, effectifs, utilisateurs, organisation financière et priorités.",
        ordered_num_id,
        "Cadrer les besoins :",
    )
    add_list_item(
        document,
        "Configurer l'école : identité, classes, année scolaire, tarifs, rôles et permissions.",
        ordered_num_id,
        "Configurer l'école :",
    )
    add_list_item(
        document,
        "Reprendre les données : importer les listes disponibles et contrôler les informations essentielles.",
        ordered_num_id,
        "Reprendre les données :",
    )
    add_list_item(
        document,
        "Former et accompagner : démarrer par les modules prioritaires puis étendre l'usage.",
        ordered_num_id,
        "Former et accompagner :",
    )

    # Page 5 — Schools
    add_page_break(document)
    add_kicker(document, "Références")
    add_title(document, "Établissements utilisateurs et déploiement en cours", 25, NAVY, after=5)
    add_subtitle(
        document,
        "MySchoolGN accompagne déjà plusieurs communautés scolaires et poursuit son déploiement auprès de nouveaux établissements.",
        after=10,
    )
    add_metric_strip(
        document,
        [
            ("4 établissements", "utilisent MySchoolGN"),
            ("1 établissement", "en contractualisation"),
            ("1 même ambition", "mieux piloter l'école"),
        ],
    )

    schools = [
        ("01", "Hadja Kanfing de Somayah", "EN UTILISATION", PALE_GREEN, GREEN),
        ("02", "Kinder School Internationale", "EN UTILISATION", PALE_GREEN, GREEN),
        ("03", "Les Écoles Naby Bakoro de Somayah", "EN UTILISATION", PALE_GREEN, GREEN),
        ("04", "Hadja Kanfing de Sonfonia", "EN UTILISATION", PALE_GREEN, GREEN),
        ("05", "Les Espoirs d'Afrique", "CONTRACTUALISATION EN COURS", PALE_AMBER, AMBER),
    ]
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [760, 5360, 3240])
    set_table_borders(table, color=BORDER, size=5)
    header = table.rows[0]
    set_repeat_table_header(header)
    labels = ["N°", "ÉTABLISSEMENT", "STATUT"]
    for idx, label in enumerate(labels):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=9.2, color=WHITE, bold=True)
    for number, school, status, status_fill, status_color in schools:
        row = table.add_row()
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, {"top": 135, "bottom": 135, "start": 120, "end": 120})
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": BORDER},
                bottom={"val": "single", "sz": 5, "color": BORDER},
                start={"val": "single", "sz": 5, "color": BORDER},
                end={"val": "single", "sz": 5, "color": BORDER},
            )
        p0 = clear_cell(row.cells[0])
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p0.add_run(number), size=9.5, color=GRAY, bold=True)
        p1 = clear_cell(row.cells[1])
        set_font(p1.add_run(school), size=10.5, color=INK, bold=True)
        p2 = clear_cell(row.cells[2])
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row.cells[2], status_fill)
        set_font(p2.add_run(status), size=8.8, color=status_color, bold=True)

    add_small_note(
        document,
        "Liste communiquée pour ce catalogue en juillet 2026. Le statut « contractualisation en cours » ne correspond pas encore à une mise en production.",
        after=8,
    )
    add_callout(
        document,
        "UNE SOLUTION, PLUSIEURS CONTEXTES",
        "Chaque établissement peut configurer son identité, ses cycles, ses tarifs, ses rôles et les modules utiles à son organisation.",
        fill=PALE_BLUE,
        accent=NAVY,
    )
    add_section_heading(document, "Ce que MySchoolGN structure au quotidien", after=5)
    add_cards_grid(
        document,
        [
            {
                "title": "VIE SCOLAIRE",
                "body": "Inscriptions, classes, élèves et familles.",
                "accent": NAVY,
                "fill": PALE_BLUE,
            },
            {
                "title": "PÉDAGOGIE",
                "body": "Notes, bulletins, classements et activités.",
                "accent": GREEN,
                "fill": PALE_GREEN,
            },
            {
                "title": "GESTION",
                "body": "Paiements, dépenses, salaires et services.",
                "accent": ORANGE,
                "fill": PALE_ORANGE,
            },
            {
                "title": "RELATION PARENTS",
                "body": "Rapports, reçus et communication ciblée.",
                "accent": BLUE,
                "fill": LIGHT_GRAY,
            },
        ],
    )

    # Page 6 — CTA
    add_page_break(document)
    p_logo2 = document.add_paragraph()
    p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo2.paragraph_format.space_before = Pt(6)
    p_logo2.paragraph_format.space_after = Pt(10)
    set_image_alt_text(
        p_logo2.add_run().add_picture(str(LOGO), width=Inches(1.45)),
        "Logo Africa Mindset Consulting",
        "Africa Mindset Consulting",
    )
    add_kicker(document, "Passons à l'étape suivante", WD_ALIGN_PARAGRAPH.CENTER, 5)
    add_title(
        document,
        "Prêt à moderniser la gestion de votre école ?",
        28,
        NAVY,
        WD_ALIGN_PARAGRAPH.CENTER,
        8,
    )
    add_subtitle(
        document,
        "Demandez une présentation de MySchoolGN et identifiez les modules les plus utiles à votre établissement.",
        WD_ALIGN_PARAGRAPH.CENTER,
        18,
    )

    add_metric_strip(
        document,
        [
            ("1. Démonstration", "Découvrir les parcours clés"),
            ("2. Cadrage", "Prioriser les besoins de l'école"),
            ("3. Déploiement", "Configurer, importer et former"),
        ],
    )

    add_section_heading(document, "Contact MySchoolGN", after=7)
    contact_table = document.add_table(rows=3, cols=2)
    set_table_geometry(contact_table, [2250, 7110])
    set_table_borders(contact_table, color=BORDER, size=5)
    contacts = [
        ("SITE WEB", "www.myschoolgn.space", "https://www.myschoolgn.space"),
        ("TÉLÉPHONE", "+224 622 61 35 59", "tel:+224622613559"),
        ("EMAIL", "contact@myschoolgn.space", "mailto:contact@myschoolgn.space"),
    ]
    for idx, (label, value, url) in enumerate(contacts):
        left = contact_table.cell(idx, 0)
        right = contact_table.cell(idx, 1)
        set_cell_shading(left, NAVY)
        set_cell_shading(right, WHITE)
        p_left = clear_cell(left)
        p_left.paragraph_format.space_after = Pt(0)
        set_font(p_left.add_run(label), size=9.3, color=WHITE, bold=True)
        p_right = clear_cell(right)
        p_right.paragraph_format.space_after = Pt(0)
        add_hyperlink(p_right, value, url, color=BLUE, underline=False)

    add_callout(
        document,
        "MYSCHOOLGN",
        "Gérez votre école avec intelligence.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_body(
        document,
        "Préscolaire · Primaire · Collège · Lycée",
        after=7,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_small_note(
        document,
        "Les fonctionnalités disponibles dépendent des modules activés et de la configuration retenue par l'établissement.",
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
